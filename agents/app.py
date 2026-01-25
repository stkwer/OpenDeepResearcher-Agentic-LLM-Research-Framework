import streamlit as st
import re
import time
from fpdf import FPDF
from docx import Document

from writer_agent import WriterAgent
from planner_agent import PlannerAgent
from searcher_agent import SearcherAgent

# =============================================
# PAGE CONFIG
# =============================================
st.set_page_config(
    page_title="OpenDeepResearcher",
    page_icon="🤖",
    layout="wide"
)

# =============================================
# FORCE DARK MODE + UI STYLING
# =============================================
st.markdown("""
<style>
html, body, [class*="css"] {
    background-color: #0b0f1a !important;
    color: #e5e7eb !important;
}

section[data-testid="stSidebar"] {
    background-color: #020617 !important;
    border-right: 1px solid #1f2937;
}

textarea, input {
    background-color: #020617 !important;
    color: #e5e7eb !important;
    border-radius: 14px !important;
    border: 1px solid #f5c77a40 !important;
}

button {
    background: linear-gradient(90deg, #f5c77a, #d4af37) !important;
    color: #1f2933 !important;
    border-radius: 14px !important;
    border: none !important;
    font-weight: 600 !important;
}

details {
    background-color: #020617 !important;
    border-radius: 12px;
    border: 1px solid #1f2937;
}
</style>
""", unsafe_allow_html=True)

# =============================================
# SESSION STATE
# =============================================
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "input_text" not in st.session_state:
    st.session_state.input_text = ""

# =============================================
# AGENTS
# =============================================
writer = WriterAgent()
planner_agent = PlannerAgent()
searcher_agent = SearcherAgent()

# =============================================
# HELPERS
# =============================================
def format_answer(text: str) -> str:
    """
    Converts markdown-like LLM output into clean HTML.
    Generic, topic-agnostic, no hard-coding.
    """

    # Convert markdown headings (###, ####, etc.) to bold
    text = re.sub(
        r"^\s*#{1,6}\s*(.+)$",
        r"<b>\1</b>",
        text,
        flags=re.MULTILINE
    )

    # Convert **bold** markdown to HTML bold
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)

    # Remove numbered lists (1., 2., etc.)
    text = re.sub(r"^\s*\d+\.\s*", "", text, flags=re.MULTILINE)

    # Remove bullet markers
    text = re.sub(r"^\s*[-•]\s*", "", text, flags=re.MULTILINE)

    # Bold standalone side headings ending with :
    text = re.sub(
        r"^(?!\s)(.{3,}?):\s*$",
        r"<b>\1:</b>",
        text,
        flags=re.MULTILINE
    )

    return text.strip()

def build_context(history, max_turns=3):
    context = ""
    for h in history[-max_turns:]:
        context += f"Previous question: {h['user']}\n"
        context += f"Previous answer: {h['bot']}\n\n"
    return context

def topic_title(text, max_words=6):
    words = text.split()
    return " ".join(words[:max_words]) + ("..." if len(words) > max_words else "")

# =============================================
# EXPORT HELPERS
# =============================================
def export_chat_pdf(chat_history):
    pdf = FPDF()
    pdf.set_auto_page_break(True, 15)
    pdf.add_page()
    pdf.set_font("Arial", size=11)

    for chat in chat_history:
        pdf.multi_cell(0, 8, f"User:\n{chat['user']}\n")
        pdf.multi_cell(0, 8, f"Assistant:\n{chat['bot']}\n")
        pdf.ln(4)

    path = "chat_export.pdf"
    pdf.output(path)
    return path

def export_chat_docx(chat_history):
    doc = Document()
    doc.add_heading("OpenDeepResearcher Chat Export", level=1)

    for chat in chat_history:
        doc.add_heading("User", level=3)
        doc.add_paragraph(chat["user"])
        doc.add_heading("Assistant", level=3)
        doc.add_paragraph(chat["bot"])

    path = "chat_export.docx"
    doc.save(path)
    return path

# =============================================
# SIDEBAR
# =============================================
st.sidebar.markdown("## 💬 Conversation Memory")

if st.sidebar.button("➕ New Chat"):
    st.session_state.chat_history = []
    st.session_state.input_text = ""
    st.rerun()

if st.sidebar.button("🗑 Clear Memory"):
    st.session_state.chat_history = []
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("### 📜 Chat History")

if st.session_state.chat_history:
    for chat in st.session_state.chat_history:
        st.sidebar.markdown(f"• {topic_title(chat['user'])}")
else:
    st.sidebar.info("No chats yet")

st.sidebar.markdown("---")
st.sidebar.markdown("### ⬇️ Export Chat")

if st.session_state.chat_history:
    if st.sidebar.button("📄 Export as PDF"):
        path = export_chat_pdf(st.session_state.chat_history)
        with open(path, "rb") as f:
            st.sidebar.download_button("Download PDF", f, file_name="chat.pdf")

    if st.sidebar.button("📝 Export as DOCX"):
        path = export_chat_docx(st.session_state.chat_history)
        with open(path, "rb") as f:
            st.sidebar.download_button("Download DOCX", f, file_name="chat.docx")

# =============================================
# HEADER
# =============================================
st.markdown("""
<h1 style="text-align:center;">🤖 OpenDeepResearcher</h1>
<p style="text-align:center; color:#9ca3af;">
Planner • Searcher • Writer • Session Memory
</p>
""", unsafe_allow_html=True)

# =============================================
# CHAT DISPLAY
# =============================================
for chat in st.session_state.chat_history:

    # USER MESSAGE
    st.markdown(
        f"<div style='margin-left:auto;background:#f5c77a;color:#1f2933;"
        f"padding:12px 16px;border-radius:18px 18px 4px 18px;max-width:70%;'>"
        f"{chat['user']}</div>",
        unsafe_allow_html=True
    )

    # PLANNER AGENT
    if chat.get("planner"):
        with st.expander("🧩 Planner Agent Questions"):
            for step in chat["planner"]:
                st.markdown(f"- {step}")

    # SEARCHER AGENT
    if chat.get("resources"):
        with st.expander("📚 Searcher Agent Sources"):
            for r in chat["resources"]:
                if isinstance(r, dict) and "title" in r and "url" in r:
                    st.markdown(f"- [{r['title']}]({r['url']})")
                else:
                    st.markdown(f"- {r}")

    # WRITER AGENT ANSWER
    clean_answer = format_answer(chat["bot"]).replace("\n", "<br>")
    st.markdown(
        f"<div style='margin-top:6px;background:#020617;color:#e5e7eb;"
        f"padding:14px 16px;border-radius:18px 18px 18px 4px;max-width:70%;"
        f"border:1px solid #f5c77a40;'>"
        f"<small>✍️ Writer Agent</small><br>{clean_answer}</div>",
        unsafe_allow_html=True
    )

# =============================================
# INPUT AREA
# =============================================
st.markdown("---")

with st.form("chat_form", clear_on_submit=False):
    col1, col2 = st.columns([8, 1])

    with col1:
        user_input = st.text_area(
            "User message",
            placeholder="Ask a question...",
            height=120,
            key="input_text",
            label_visibility="collapsed"
        )

    with col2:
        st.markdown("<div style='height:46px'></div>", unsafe_allow_html=True)
        send = st.form_submit_button("Send 🚀")

# =============================================
# MAIN LOGIC
# =============================================
if send and user_input.strip():
    context = build_context(st.session_state.chat_history)

    with st.spinner("✍️ Writer is thinking..."):
        try:
            answer = writer.answer(context + user_input)
        except Exception:
            time.sleep(2)
            answer = "⚠️ Response timeout. Please try a shorter query."

    try:
        planner_steps = planner_agent.plan(user_input).get("steps", [])
    except Exception:
        planner_steps = []

    try:
        resources = searcher_agent.answer(user_input)[:5]
    except Exception:
        resources = []

    st.session_state.chat_history.append({
        "user": user_input,
        "bot": answer,
        "planner": planner_steps,
        "resources": resources
    })

    st.rerun()
